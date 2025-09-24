
from __future__ import annotations
import io
import re
from typing import Dict, Any, List, Optional

from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docxcompose.composer import Composer


def _doc_to_bytes(doc: Document) -> bytes:
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_default_template() -> bytes:
    """Create a minimal .docx template with sane defaults and return its bytes."""
    doc = Document()
    # Page setup
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(15)

    # Header / Footer
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = "New Layout"
    header_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = "Confidential"
    footer_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Basic style defaults (Normal + Heading 1..3)
    styles = doc.styles
    try:
        styles['Normal'].font.name = 'Calibri'
        styles['Normal'].font.size = Pt(11)
        styles['Heading 1'].font.size = Pt(16)
        styles['Heading 2'].font.size = Pt(14)
        styles['Heading 3'].font.size = Pt(12)
    except KeyError:
        pass

    return _doc_to_bytes(doc)


def compose_into_template(src_docx_bytes: bytes, template_docx_bytes: bytes) -> bytes:
    """Append the legacy document into the template, preserving content (images, tables, lists)."""
    base = Document(io.BytesIO(template_docx_bytes))
    comp = Composer(base)
    src = Document(io.BytesIO(src_docx_bytes))
    comp.append(src)
    bio = io.BytesIO()
    comp.save(bio)
    return bio.getvalue()


def apply_style_mappings(docx_bytes: bytes, mapping: Dict[str, str]) -> bytes:
    """Map paragraph style names to target names (best-effort)."""
    doc = Document(io.BytesIO(docx_bytes))
    for p in doc.paragraphs:
        try:
            if p.style and p.style.name in mapping:
                p.style = doc.styles[mapping[p.style.name]]
        except Exception:
            # If style missing or mapping invalid, skip
            pass
    return _doc_to_bytes(doc)


def apply_find_replace(docx_bytes: bytes, rules: List[Dict[str, Any]]) -> bytes:
    """Apply regex find/replace to paragraph texts (best-effort, may simplify run formatting)."""
    doc = Document(io.BytesIO(docx_bytes))

    def replace_in_paragraph(p, pattern: str, repl: str):
        # Join runs to a single string
        text = "".join(run.text for run in p.runs) or p.text
        new_text = re.sub(pattern, repl, text, flags=re.MULTILINE)
        if new_text != text:
            # Clear runs and set one run with new text (may lose mixed formatting)
            for _ in range(len(p.runs)):
                p.runs[0].clear()
                p.runs[0]._r.getparent().remove(p.runs[0]._r)
            p.add_run(new_text)

    for p in doc.paragraphs:
        for rule in rules:
            pattern = rule.get("pattern")
            repl = rule.get("replace", "")
            if pattern:
                replace_in_paragraph(p, pattern, repl)

    # Header / footer replacements
    for section in doc.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                for rule in rules:
                    pattern = rule.get("pattern")
                    repl = rule.get("replace", "")
                    if pattern:
                        replace_in_paragraph(p, pattern, repl)

    return _doc_to_bytes(doc)


def set_header_footer(docx_bytes: bytes, cfg: Dict[str, Any]) -> bytes:
    """Set header/footer text and optional page numbers."""
    doc = Document(io.BytesIO(docx_bytes))
    header_text = cfg.get("header_text")
    footer_text = cfg.get("footer_text")
    include_page_numbers = bool(cfg.get("include_page_numbers", False))

    for section in doc.sections:
        if header_text is not None:
            hdr = section.header
            if not hdr.paragraphs:
                hdr.add_paragraph("")
            p = hdr.paragraphs[0]
            p.text = header_text

        if footer_text is not None or include_page_numbers:
            ftr = section.footer
            p = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph("")
            parts = []
            if footer_text:
                parts.append(footer_text)
            if include_page_numbers:
                if parts:
                    parts.append(" — ")
                # Insert PAGE field
                parts.append("Page ")
                run = p.add_run("")
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = "PAGE"
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'separate')
                fldChar3 = OxmlElement('w:fldChar')
                fldChar3.set(qn('w:fldCharType'), 'end')
                r_element = run._r
                r_element.append(fldChar1)
                r_element.append(instrText)
                r_element.append(fldChar2)
                # place holder text that Word will update
                t = OxmlElement('w:t')
                t.text = "1"
                r_element.append(t)
                r_element.append(fldChar3)
            # Put initial text at the start
            if parts:
                if p.text:
                    p.text = "".join(parts) + p.text
                else:
                    p.text = "".join(parts)

    return _doc_to_bytes(doc)


def apply_page_setup(docx_bytes: bytes, cfg: Dict[str, Any]) -> bytes:
    """Apply page orientation and margins."""
    doc = Document(io.BytesIO(docx_bytes))
    orientation = str(cfg.get("orientation", "portrait")).lower()
    margins = cfg.get("margins_mm", {})
    for section in doc.sections:
        if orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        else:
            section.orientation = WD_ORIENT.PORTRAIT
        # Margins
        for side, attr in (("top", "top_margin"), ("bottom", "bottom_margin"),
                           ("left", "left_margin"), ("right", "right_margin")):
            if side in margins:
                setattr(section, attr, Mm(float(margins[side])))
    return _doc_to_bytes(doc)
